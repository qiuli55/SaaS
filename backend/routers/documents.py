import os
import re
import json
import httpx
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from typing import Optional
from datetime import datetime
import io

from database import get_db
from models import Document, Case, User
from schemas import DocumentGenerate, DocumentInfo
from auth import get_current_user

router = APIRouter(tags=["文书"])

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEEPSEEK_BASE_URL = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com/v1")

# 文书类型与格式配置
DOC_TYPES = ["民事起诉状", "民事答辩状", "律师函", "代理词", "法律意见书", "上诉状", "再审申请书", "催收函"]

# 核心法条库（用于校验）
LAW_ARTICLES = {
    "民法典": {
        "577": "当事人一方不履行合同义务或者履行合同义务不符合约定的，应当承担继续履行、采取补救措施或者赔偿损失等违约责任。",
        "578": "当事人一方明确表示或者以自己的行为表明不履行合同义务的，对方可以在履行期限届满前请求其承担违约责任。",
        "579": "当事人一方未支付价款、报酬、租金、利息，或者不履行其他金钱债务的，对方可以请求其支付。",
        "580": "当事人一方不履行非金钱债务或者履行非金钱债务不符合约定的，对方可以请求履行。",
        "583": "当事人一方不履行合同义务或者履行合同义务不符合约定的，在履行义务或者采取补救措施后，对方还有其他损失的，应当赔偿损失。",
        "584": "当事人一方不履行合同义务或者履行合同义务不符合约定，造成对方损失的，损失赔偿额应当相当于因违约所造成的损失。",
        "585": "当事人可以约定一方违约时应当根据违约情况向对方支付一定数额的违约金。",
        "595": "买卖合同是出卖人转移标的物的所有权于买受人，买受人支付价款的合同。",
        "626": "买受人应当按照约定的数额和支付方式支付价款。",
        "628": "买受人应当按照约定的时间支付价款。",
        "667": "借款合同是借款人向贷款人借款，到期返还借款并支付利息的合同。",
        "675": "借款人应当按照约定的期限返还借款。",
        "676": "借款人未按照约定的期限返还借款的，应当按照约定或者国家有关规定支付逾期利息。",
        "1079": "夫妻一方要求离婚的，可以由有关组织进行调解或者直接向人民法院提起离婚诉讼。",
        "1084": "父母与子女间的关系，不因父母离婚而消除。离婚后，子女无论由父或者母直接抚养，仍是父母双方的子女。",
        "1087": "离婚时，夫妻的共同财产由双方协议处理；协议不成的，由人民法院根据财产的具体情况判决。",
    },
    "民事诉讼法": {
        "122": "起诉必须符合下列条件：（一）原告是与本案有直接利害关系的公民、法人和其他组织；（二）有明确的被告；（三）有具体的诉讼请求和事实、理由；（四）属于人民法院受理民事诉讼的范围和受诉人民法院管辖。",
        "123": "人民法院应当保障当事人依照法律规定享有的起诉权利。对符合本法第一百二十二条的起诉，必须受理。",
        "264": "被告经传票传唤，无正当理由拒不到庭的，或者未经法庭许可中途退庭的，可以缺席判决。",
    },
}

SYSTEM_PROMPT = """你是一名有10年执业经验的中国律师，擅长撰写法律文书。
请根据用户提供的案件信息，撰写一份符合中国法院要求的{doc_type}。

格式要求：
1. 标题：居中，加粗
2. 当事人信息：分段列出，包含姓名/名称、住址/住所地、法定代表人等信息
3. 诉讼请求（如适用）：分项列出，每项以"1." "2."开头
4. 事实与理由：分段叙述，逻辑清晰，时间线清楚
5. 此致：右对齐，后接法院名称
6. 落款：右对齐，具状人/申请人，日期

内容要求：
1. 法条引用必须准确，只能引用真实存在的中国法律
2. 不得虚构事实、证据或当事人未提供的信息
3. 语言严谨规范，使用法律术语
4. 必须明列具体法条，不得使用"根据相关法律规定"等模糊表述
5. 引用法条时使用格式：《中华人民共和国民法典》第XXX条

输出格式：
首先输出完整的法律文书正文。
然后在末尾用单独一行"===法条引用==="开始，逐行列出所有引用的法条，每行一个法条。"""


def verify_articles(text: str) -> list[dict]:
    """校验 AI 输出的法条引用是否真实存在"""
    results = []
    # 匹配 "《xxx》第X条" 格式
    pattern = r'《(.+?)》第(\d+)条'
    matches = re.findall(pattern, text)
    seen = set()
    for law_name, article_num in matches:
        key = f"{law_name}_{article_num}"
        if key in seen:
            continue
        seen.add(key)

        found = False
        detail = ""
        for short_name, articles in LAW_ARTICLES.items():
            if short_name in law_name:
                if article_num in articles:
                    found = True
                    detail = articles[article_num]
                    break

        results.append({
            "law": f"《{law_name}》",
            "article": f"第{article_num}条",
            "verified": found,
            "detail": detail if found else "未在法条库中找到，请人工核实",
        })
    return results


@router.post("/api/documents/generate")
async def generate_document(
    req: DocumentGenerate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    if not DEEPSEEK_API_KEY or DEEPSEEK_API_KEY == "your-deepseek-api-key-here":
        raise HTTPException(status_code=500, detail="请先配置 DeepSeek API Key")

    # 获取案件信息
    case = db.query(Case).filter(
        Case.id == req.case_id, Case.user_id == current_user.id
    ).first()
    if not case:
        raise HTTPException(status_code=404, detail="案件不存在")

    if req.doc_type not in DOC_TYPES:
        raise HTTPException(status_code=400, detail=f"不支持的文书类型：{req.doc_type}")

    # 构建提示词
    user_prompt = f"""请根据以下信息生成一份{req.doc_type}：

案由：{case.case_type}
原告/申请人：{case.plaintiff}
被告/被申请人：{case.defendant}
标的额：{case.subject_amount or '未填写'}元
诉讼请求：{req.claims or '请根据案件事实合理推断'}
案件事实：{req.facts or case.description or '请根据案由和当事人信息合理撰写'}

请严格遵守格式要求，写出完整的法律文书。"""

    sp = SYSTEM_PROMPT.replace("{doc_type}", req.doc_type)

    # 调 DeepSeek API
    async with httpx.AsyncClient(timeout=120.0) as client:
        try:
            resp = await client.post(
                f"{DEEPSEEK_BASE_URL}/chat/completions",
                headers={
                    "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "deepseek-chat",
                    "messages": [
                        {"role": "system", "content": sp},
                        {"role": "user", "content": user_prompt},
                    ],
                    "temperature": 0.3,
                    "max_tokens": 8192,
                },
            )
            resp.raise_for_status()
            result = resp.json()
        except httpx.HTTPStatusError as e:
            raise HTTPException(status_code=502, detail=f"AI 服务请求失败: {e.response.text}")
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"AI 服务异常: {str(e)}")

    ai_text = result["choices"][0]["message"]["content"]

    # 分离文书正文和法条引用
    final_content = ai_text
    if "===法条引用===" in ai_text:
        parts = ai_text.split("===法条引用===", 1)
        final_content = parts[0].strip()
        articles_text = parts[1].strip() if len(parts) > 1 else ""
    else:
        articles_text = ai_text

    # 法条校验
    verified_articles = verify_articles(articles_text)

    # 计算版本号
    existing_versions = db.query(Document).filter(
        Document.case_id == req.case_id,
        Document.doc_type == req.doc_type,
    ).count()

    # 保存文书
    doc = Document(
        case_id=req.case_id,
        user_id=current_user.id,
        doc_type=req.doc_type,
        version=existing_versions + 1,
        form_data={
            "claims": req.claims,
            "facts": req.facts,
        },
        ai_raw_text=ai_text,
        final_content=final_content,
        verified_articles=verified_articles,
        status="已完成",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return {
        "code": 0,
        "message": "生成成功",
        "data": {
            "id": doc.id,
            "case_id": doc.case_id,
            "doc_type": doc.doc_type,
            "version": doc.version,
            "final_content": doc.final_content,
            "verified_articles": doc.verified_articles,
            "status": doc.status,
            "created_at": doc.created_at.isoformat(),
        },
    }


@router.get("/api/cases/{case_id}/documents")
def list_documents(
    case_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    docs = db.query(Document).filter(
        Document.case_id == case_id,
        Document.user_id == current_user.id,
    ).order_by(Document.created_at.desc()).all()

    case = db.query(Case).filter(Case.id == case_id).first()
    case_name = f"{case.plaintiff}{case.case_type}" if case else ""

    items = []
    for d in docs:
        items.append(DocumentInfo(
            id=d.id,
            case_id=d.case_id,
            doc_type=d.doc_type,
            version=d.version,
            form_data=d.form_data,
            final_content=d.final_content,
            verified_articles=d.verified_articles,
            status=d.status,
            created_at=d.created_at,
            case_name=case_name,
            plaintiff=case.plaintiff if case else "",
            defendant=case.defendant if case else "",
        ))

    return {"code": 0, "data": [item.model_dump() for item in items]}


@router.get("/api/documents/{doc_id}")
def get_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = db.query(Document).filter(
        Document.id == doc_id,
        Document.user_id == current_user.id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文书不存在")

    case = db.query(Case).filter(Case.id == doc.case_id).first()

    return {
        "code": 0,
        "data": DocumentInfo(
            id=doc.id,
            case_id=doc.case_id,
            doc_type=doc.doc_type,
            version=doc.version,
            form_data=doc.form_data,
            final_content=doc.final_content,
            verified_articles=doc.verified_articles,
            status=doc.status,
            created_at=doc.created_at,
            case_name=f"{case.plaintiff}{case.case_type}" if case else "",
            plaintiff=case.plaintiff if case else "",
            defendant=case.defendant if case else "",
        ).model_dump(),
    }


@router.get("/api/documents/{doc_id}/download/docx")
def download_docx(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = db.query(Document).filter(
        Document.id == doc_id,
        Document.user_id == current_user.id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文书不存在")

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        word_doc = DocxDocument()

        # 设置默认字体
        style = word_doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(14)

        lines = doc.final_content.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            p = word_doc.add_paragraph()

            # 标题行
            if i == 0 and (doc.doc_type in line or "起诉状" in line or "答辩状" in line or
                           "律师函" in line or "代理词" in line or "意见书" in line or
                           "上诉状" in line or "申请书" in line or "函" in line):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line)
                run.font.size = Pt(22)
                run.font.name = '宋体'
                run.bold = True
            elif "此致" in line:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(line)
                run.font.size = Pt(14)
                run.font.name = '宋体'
            elif "法院" in line and ("此致" in lines[max(0, i-1)] or
                                     any("此致" in lines[j] for j in range(max(0, i-3), i))):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(line)
                run.font.size = Pt(14)
                run.font.name = '宋体'
            elif "起诉人" in line or "具状人" in line or "申请人" in line or "答辩人" in line or \
                 "年" in line and "月" in line and "日" in line and len(line) < 20:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(line)
                run.font.size = Pt(14)
                run.font.name = '宋体'
            else:
                run = p.add_run(line)
                run.font.size = Pt(14)
                run.font.name = '宋体'

            p.paragraph_format.line_spacing = Pt(28)
            p.paragraph_format.space_after = Pt(0)

        buf = io.BytesIO()
        word_doc.save(buf)
        buf.seek(0)

        filename = f"{doc.doc_type}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename.encode('utf-8').decode('latin-1')}"},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出 Word 失败: {str(e)}")


@router.get("/api/documents/{doc_id}/download/pdf")
def download_pdf(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = db.query(Document).filter(
        Document.id == doc_id,
        Document.user_id == current_user.id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文书不存在")

    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=25)

        # 尝试使用中文字体
        chinese_font = None
        font_paths = [
            "C:/Windows/Fonts/simsun.ttc",
            "C:/Windows/Fonts/simfang.ttf",
            "C:/Windows/Fonts/msyh.ttc",
        ]
        for fp in font_paths:
            if os.path.exists(fp):
                pdf.add_font("CN", "", fp, uni=True)
                pdf.add_font("CN", "B", fp, uni=True)
                chinese_font = "CN"
                break

        if not chinese_font:
            # 兜底：返回纯文本 PDF
            pdf.set_font("Helvetica", "", 12)
            lines = doc.final_content.split("\n")
            for line in lines:
                line = line.strip()
                safe_line = line.encode("latin-1", errors="replace").decode("latin-1")
                if not safe_line.strip():
                    pdf.ln(5)
                else:
                    pdf.cell(0, 8, safe_line, ln=True)
        else:
            pdf.set_font(chinese_font, "", 14)
            lines = doc.final_content.split("\n")
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    pdf.ln(5)
                    continue

                # 标题
                if i == 0 and (doc.doc_type in line or any(kw in line for kw in
                    ["起诉状", "答辩状", "律师函", "代理词", "意见书", "上诉状", "申请书", "函"])):
                    pdf.set_font(chinese_font, "B", 18)
                    pdf.cell(0, 14, line, ln=True, align="C")
                    pdf.set_font(chinese_font, "", 14)
                # 此致 / 落款
                elif "此致" in line and len(line) <= 5:
                    pdf.cell(0, 10, line, ln=True, align="R")
                elif "法院" in line:
                    pdf.cell(0, 10, line, ln=True, align="L")
                    pdf.ln(5)
                elif "具状人" in line or "起诉人" in line or "申请人" in line or "答辩人" in line:
                    pdf.cell(0, 14, line, ln=True, align="R")
                elif len(line) < 20 and "年" in line and "月" in line and "日" in line:
                    pdf.cell(0, 10, line, ln=True, align="R")
                else:
                    pdf.multi_cell(0, 8, line)

        buf = io.BytesIO()
        pdf.output(buf)
        buf.seek(0)

        filename = f"{doc.doc_type}.pdf"
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename.encode('utf-8').decode('latin-1')}"},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出 PDF 失败: {str(e)}")


@router.get("/api/documents/history")
def document_history(
    keyword: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    query = db.query(Document).filter(Document.user_id == current_user.id)
    if keyword:
        query = query.filter(
            (Document.doc_type.contains(keyword)) |
            (Document.final_content.contains(keyword))
        )

    total = query.count()
    docs = query.order_by(Document.created_at.desc()).offset(
        (page - 1) * page_size
    ).limit(page_size).all()

    items = []
    for d in docs:
        case = db.query(Case).filter(Case.id == d.case_id).first()
        items.append(DocumentInfo(
            id=d.id,
            case_id=d.case_id,
            doc_type=d.doc_type,
            version=d.version,
            form_data=d.form_data,
            final_content=d.final_content[:200] if d.final_content else "",
            verified_articles=d.verified_articles,
            status=d.status,
            created_at=d.created_at,
            case_name=f"{case.plaintiff}{case.case_type}" if case else "",
            plaintiff=case.plaintiff if case else "",
            defendant=case.defendant if case else "",
        ))

    return {"code": 0, "data": {
        "total": total,
        "items": [item.model_dump() for item in items],
        "page": page,
        "page_size": page_size,
    }}
